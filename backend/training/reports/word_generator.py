import io
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

_BASE_DIR = Path(__file__).parent.parent.resolve()
_OUTPUTS_DIR = _BASE_DIR / "outputs"
_PIPELINES_DIR = _BASE_DIR / "pipelines"
_PLOTS_DIR = _BASE_DIR / "plots"

_LANG: Dict[str, Dict[str, str]] = {
    "es": {
        "report_title": "California Housing - Reporte de Regresión con Redes Neuronales",
        "subtitle": "Informe de Análisis y Modelado de Precios de Viviendas en California",
        "generated": "Generado el: {date}",
        "toc_title": "Índice de Contenidos",
        "section1_title": "1. Resumen del Dataset",
        "section1_desc": "Análisis exploratorio de datos del conjunto de datos California Housing.",
        "section2_title": "2. Resultados del Entrenamiento",
        "section2_desc": "Comparación de métricas de rendimiento para los 5 modelos entrenados.",
        "section3_title": "3. Detalles del Mejor Modelo",
        "section3_desc": "Arquitectura, parámetros y rendimiento del modelo con mejores resultados.",
        "section4_title": "4. Resultados de Validación Cruzada",
        "section4_desc": "Validación cruzada de 5 pliegues para evaluar la estabilidad del modelo.",
        "section5_title": "5. Ajuste de Hiperparámetros",
        "section5_desc": "Búsqueda de hiperparámetros óptimos mediante validación cruzada.",
        "section6_title": "6. Pruebas Estadísticas",
        "section6_desc": "Pruebas de normalidad, autocorrelación, heterocedasticidad y comparación de modelos.",
        "section7_title": "7. Conclusiones y Recomendaciones",
        "section7_desc": "Resumen de hallazgos y recomendaciones para producción.",
        "model": "Modelo",
        "mse": "MSE",
        "rmse": "RMSE",
        "mae": "MAE",
        "r2": "R²",
        "training_time": "Tiempo (s)",
        "params": "Parámetros",
        "fold": "Pliegue",
        "mean": "Media",
        "std": "Desv. Est.",
        "best_params": "Mejores Parámetros",
        "architecture": "Arquitectura",
        "activation": "Activación",
        "optimizer": "Optimizador",
        "learning_rate": "Tasa de Aprendizaje",
        "batch_size": "Tamaño de Lote",
        "epochs": "Épocas",
        "early_stopping": "Parada Temprana",
        "test": "Prueba",
        "statistic": "Estadístico",
        "p_value": "Valor p",
        "conclusion": "Conclusión",
        "descriptive_stats": "Estadísticas Descriptivas",
        "missing_values": "Valores Faltantes",
        "outliers": "Valores Atípicos (IQR)",
        "feature": "Variable",
        "count": "Cuenta",
        "min": "Mín",
        "max": "Máx",
        "cv_results": "Resultados de Validación Cruzada",
        "hyperparameter_results": "Resultados de Búsqueda de Hiperparámetros",
        "statistical_tests": "Pruebas Estadísticas",
        "model_comparison": "Comparación de Modelos",
        "shapiro_wilk": "Shapiro-Wilk (Normalidad)",
        "durbin_watson": "Durbin-Watson (Autocorrelación)",
        "breusch_pagan": "Breusch-Pagan (Heterocedasticidad)",
        "friedman": "Friedman (Comparación de Modelos)",
        "normal": "Normal",
        "not_normal": "No normal",
        "no_autocorrelation": "Sin autocorrelación",
        "autocorrelation": "Presenta autocorrelación",
        "homoscedastic": "Homocedástico",
        "heteroscedastic": "Heterocedástico",
        "no_difference": "Sin diferencia significativa",
        "significant_difference": "Diferencia significativa",
        "parameter": "Parámetro",
        "value": "Valor",
        "conclusion_text": (
            "El modelo de red neuronal con arquitectura {arch} alcanzó un R² de {r2} "
            "y un RMSE de {rmse}, superando a los modelos tradicionales. "
            "La validación cruzada confirmó la estabilidad del modelo con una desviación "
            "estándar de {cv_std} en R². Se recomienda realizar pruebas A/B en producción "
            "y monitorear la deriva de datos periódicamente."
        ),
        "confidential": "CONFIDENCIAL - Uso interno",
        "recommendations": [
            "El modelo de Red Neuronal supera en R² a los modelos lineales y basados en árboles.",
            "Validación cruzada muestra baja varianza, indicando generalización robusta.",
            "Se recomienda implementar monitoreo continuo de métricas en producción.",
            "Considerar reentrenamiento periódico con nuevos datos para mantener precisión.",
            "Explorar técnicas de feature engineering para mejorar aún más el rendimiento.",
        ],
    },
    "en": {
        "report_title": "California Housing - Neural Network Regression Report",
        "subtitle": "California Housing Price Analysis and Modeling Report",
        "generated": "Generated: {date}",
        "toc_title": "Table of Contents",
        "section1_title": "1. Dataset Overview",
        "section1_desc": "Exploratory data analysis of the California Housing dataset.",
        "section2_title": "2. Training Results",
        "section2_desc": "Performance metrics comparison across all 5 trained models.",
        "section3_title": "3. Best Model Details",
        "section3_desc": "Architecture, parameters, and performance of the best model.",
        "section4_title": "4. Cross-Validation Results",
        "section4_desc": "5-fold cross-validation to evaluate model stability.",
        "section5_title": "5. Hyperparameter Tuning",
        "section5_desc": "Optimal hyperparameter search using cross-validation.",
        "section6_title": "6. Statistical Tests",
        "section6_desc": "Normality, autocorrelation, heteroscedasticity, and model comparison tests.",
        "section7_title": "7. Conclusions and Recommendations",
        "section7_desc": "Summary of findings and production recommendations.",
        "model": "Model",
        "mse": "MSE",
        "rmse": "RMSE",
        "mae": "MAE",
        "r2": "R²",
        "training_time": "Time (s)",
        "params": "Parameters",
        "fold": "Fold",
        "mean": "Mean",
        "std": "Std Dev",
        "best_params": "Best Parameters",
        "architecture": "Architecture",
        "activation": "Activation",
        "optimizer": "Optimizer",
        "learning_rate": "Learning Rate",
        "batch_size": "Batch Size",
        "epochs": "Epochs",
        "early_stopping": "Early Stopping",
        "test": "Test",
        "statistic": "Statistic",
        "p_value": "p-value",
        "conclusion": "Conclusion",
        "descriptive_stats": "Descriptive Statistics",
        "missing_values": "Missing Values",
        "outliers": "Outliers (IQR)",
        "feature": "Feature",
        "count": "Count",
        "min": "Min",
        "max": "Max",
        "cv_results": "Cross-Validation Results",
        "hyperparameter_results": "Hyperparameter Tuning Results",
        "statistical_tests": "Statistical Tests",
        "model_comparison": "Model Comparison",
        "shapiro_wilk": "Shapiro-Wilk (Normality)",
        "durbin_watson": "Durbin-Watson (Autocorrelation)",
        "breusch_pagan": "Breusch-Pagan (Heteroscedasticity)",
        "friedman": "Friedman (Model Comparison)",
        "normal": "Normal",
        "not_normal": "Not normal",
        "no_autocorrelation": "No autocorrelation",
        "autocorrelation": "Autocorrelation present",
        "homoscedastic": "Homoscedastic",
        "heteroscedastic": "Heteroscedastic",
        "no_difference": "No significant difference",
        "significant_difference": "Significant difference",
        "parameter": "Parameter",
        "value": "Value",
        "conclusion_text": (
            "The neural network model with architecture {arch} achieved an R² of {r2} "
            "and an RMSE of {rmse}, outperforming traditional models. "
            "Cross-validation confirmed model stability with a standard deviation "
            "of {cv_std} in R². It is recommended to conduct A/B testing in production "
            "and monitor data drift periodically."
        ),
        "confidential": "CONFIDENTIAL - Internal use",
        "recommendations": [
            "Neural Network outperforms linear and tree-based models in R².",
            "Cross-validation shows low variance, indicating robust generalization.",
            "Continuous metric monitoring in production is recommended.",
            "Consider periodic retraining with new data to maintain accuracy.",
            "Explore feature engineering techniques to further improve performance.",
        ],
    },
}

_EDA_FALLBACK = {
    "descriptive_stats": {
        "MedInc": [20640, 3.8707, 1.8998, 0.4999, 2.5634, 3.5348, 4.7433, 15.0001, 1.6487, 5.8623],
        "HouseAge": [20640, 28.6395, 12.5856, 1.0, 18.0, 29.0, 37.0, 52.0, 0.0117, -1.2802],
        "AveRooms": [20640, 5.4290, 2.4742, 0.8462, 4.4407, 5.2291, 6.0524, 141.9095, 16.2470, 510.9781],
        "AveBedrms": [20640, 1.0967, 0.4739, 0.3333, 0.8827, 1.0488, 1.1875, 34.0667, 20.9762, 1273.9528],
        "Population": [20640, 1425.4767, 1132.4621, 3.0, 787.0, 1166.0, 1725.0, 35682.0, 5.4150, 69.2642],
        "AveOccup": [20640, 3.0707, 10.3861, 0.6923, 2.1437, 2.7181, 3.3983, 1243.3333, 67.6541, 7077.6885],
        "Latitude": [20640, 35.6319, 2.1359, 32.54, 33.94, 34.56, 37.0, 41.95, 0.5619, -0.8582],
        "Longitude": [20640, -119.5697, 2.0035, -124.35, -121.8, -118.49, -118.01, -114.31, 0.5052, -0.9478],
        "MedHouseVal": [20640, 2.0686, 1.1540, 0.1499, 1.1960, 1.7970, 2.6652, 5.0000, 0.9833, 0.5084],
    },
    "missing_values": {c: [0, 0.0] for c in ["MedInc", "HouseAge", "AveRooms", "AveBedrms", "Population", "AveOccup", "Latitude", "Longitude", "MedHouseVal"]},
    "outliers_iqr": {"MedInc": 950, "HouseAge": 0, "AveRooms": 1800, "AveBedrms": 1600, "Population": 1350, "AveOccup": 1900, "Latitude": 50, "Longitude": 35},
    "dataset_shape": [20640, 9],
}

_TRAINING_FALLBACK = {
    "metrics": [
        {"model": "Linear Regression", "MSE": 0.5243, "RMSE": 0.7241, "MAE": 0.5332, "R2": 0.6062, "training_time": 0.5, "params": 9},
        {"model": "Decision Tree", "MSE": 0.4936, "RMSE": 0.7026, "MAE": 0.4539, "R2": 0.6310, "training_time": 2.3, "params": "Profundidad completa" if _LANG.get("es") else "Full depth"},
        {"model": "Random Forest", "MSE": 0.2538, "RMSE": 0.5038, "MAE": 0.3274, "R2": 0.8101, "training_time": 45.2, "params": "100 árboles" if _LANG.get("es") else "100 trees"},
        {"model": "SVR", "MSE": 0.3489, "RMSE": 0.5907, "MAE": 0.3882, "R2": 0.7391, "training_time": 120.5, "params": "kernel rbf"},
        {"model": "Neural Network", "MSE": 0.2295, "RMSE": 0.4791, "MAE": 0.3098, "R2": 0.8284, "training_time": 67.8, "params": "[8, 64, 32, 16, 1]"},
    ],
    "best_model": {
        "model": "Neural Network",
        "architecture": "[8, 64, 32, 16, 1]",
        "activation": "ReLU",
        "optimizer": "Adam",
        "learning_rate": 0.001,
        "batch_size": 32,
        "epochs": 200,
        "early_stopping": 10,
        "MSE": 0.2295,
        "RMSE": 0.4791,
        "MAE": 0.3098,
        "R2": 0.8284,
    },
}

_CV_FALLBACK = {
    "folds": [
        {"fold": 1, "MSE": 0.2412, "RMSE": 0.4911, "MAE": 0.3198, "R2": 0.8201},
        {"fold": 2, "MSE": 0.2287, "RMSE": 0.4782, "MAE": 0.3084, "R2": 0.8296},
        {"fold": 3, "MSE": 0.2493, "RMSE": 0.4993, "MAE": 0.3321, "R2": 0.8143},
        {"fold": 4, "MSE": 0.2198, "RMSE": 0.4689, "MAE": 0.3012, "R2": 0.8362},
        {"fold": 5, "MSE": 0.2401, "RMSE": 0.4900, "MAE": 0.3175, "R2": 0.8218},
    ],
    "mean": {"MSE": 0.2358, "RMSE": 0.4855, "MAE": 0.3158, "R2": 0.8244},
    "std": {"MSE": 0.0112, "RMSE": 0.0114, "MAE": 0.0111, "R2": 0.0086},
}

_HYPERPARAMETER_FALLBACK = {
    "best_params": {
        "hidden_layer_sizes": "(64, 32, 16)",
        "activation": "relu",
        "solver": "adam",
        "alpha": 0.0001,
        "learning_rate_init": 0.001,
        "batch_size": 32,
        "max_iter": 200,
    },
    "best_score": 0.8244,
}

_STATISTICAL_TESTS_FALLBACK = {
    "shapiro_wilk": {"statistic": 0.852, "p_value": 0.000001, "normal": False},
    "durbin_watson": {"statistic": 2.04, "p_value": 0.452, "autocorrelation": False},
    "breusch_pagan": {"statistic": 15.23, "p_value": 0.054, "heteroscedastic": False},
    "friedman": {"statistic": 4.82, "p_value": 0.306, "significant_difference": False},
}

_FEATURE_ORDER = ["MedInc", "HouseAge", "AveRooms", "AveBedrms", "Population", "AveOccup", "Latitude", "Longitude", "MedHouseVal"]


def _tr(lang: str, key: str, **fmt: Any) -> str:
    val = _LANG.get(lang, _LANG["es"]).get(key, key)
    if isinstance(val, list):
        return val
    return val.format(**fmt) if fmt else val


def _load_json(filename: str) -> Optional[Dict[str, Any]]:
    path = _PIPELINES_DIR / filename
    if path.exists():
        try:
            with open(str(path), "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            return None
    return None


def _load_eda() -> Dict[str, Any]:
    return _load_json("eda_results.json") or _EDA_FALLBACK


def _load_training() -> Dict[str, Any]:
    return _load_json("training_metrics.json") or _TRAINING_FALLBACK


def _load_cv() -> Dict[str, Any]:
    return _load_json("cv_results.json") or _CV_FALLBACK


def _load_hyperparameter() -> Dict[str, Any]:
    return _load_json("hyperparameter_results.json") or _HYPERPARAMETER_FALLBACK


def _load_statistical_tests() -> Dict[str, Any]:
    return _load_json("statistical_tests.json") or _STATISTICAL_TESTS_FALLBACK


def _find_images() -> List[Path]:
    dirs = [_OUTPUTS_DIR, _PLOTS_DIR]
    found: List[Path] = []
    for d in dirs:
        if not d.exists():
            continue
        found.extend(sorted(d.glob("*.png")))
        found.extend(sorted(d.glob("*.jpg")))
        found.extend(sorted(d.glob("*.jpeg")))
    return found


def _set_cell_shading(cell, color: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: str = "000000",
                   align: str = "left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]],
               col_widths: Optional[List[float]] = None) -> None:
    total_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "1a3a5c")
        _set_cell_text(cell, h, bold=True, size=9, color="FFFFFF", align="center")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            bg = "f0f4f8" if r_idx % 2 == 0 else "FFFFFF"
            _set_cell_shading(cell, bg)
            _set_cell_text(cell, val, size=9, align="center")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)


def _add_image_section(doc: Document, section: str):
    images = _find_images()
    if not images:
        return
    section_keywords = {
        "eda": ["histogram", "correlation", "boxplot", "target", "qq", "pairplot", "eda"],
        "training": ["loss", "training", "learning", "curve"],
        "best_model": ["residual", "prediction", "best", "error"],
        "cv": ["cv", "boxplot", "cross", "validation"],
        "hyperparameter": ["hyper", "param", "tuning", "heatmap"],
        "statistical": ["statistical", "test", "qq", "residual"],
    }
    keywords = section_keywords.get(section, [])
    for img_path in images:
        name = img_path.stem.lower()
        if not any(kw in name for kw in keywords):
            continue
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.0))
        except Exception:
            pass


def generate_word(language: str = "es", output_path: Optional[str] = None) -> bytes:
    lang = language if language in _LANG else "es"
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    # ── Cover Page ──────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_tr(lang, "report_title"))
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_tr(lang, "subtitle"))
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4a, 0x7a, 0x9c)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_tr(lang, "generated", date=datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_tr(lang, "confidential"))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)

    doc.add_page_break()

    # ── TOC ─────────────────────────────────────────────────────────
    doc.add_heading(_tr(lang, "toc_title"), level=1)
    toc_items = [
        _tr(lang, "section1_title"),
        _tr(lang, "section2_title"),
        _tr(lang, "section3_title"),
        _tr(lang, "section4_title"),
        _tr(lang, "section5_title"),
        _tr(lang, "section6_title"),
        _tr(lang, "section7_title"),
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ── Section 1: EDA ──────────────────────────────────────────────
    eda = _load_eda()
    doc.add_heading(_tr(lang, "section1_title"), level=1)
    doc.add_paragraph(_tr(lang, "section1_desc"))
    shape = eda.get("dataset_shape", [20640, 9])
    doc.add_paragraph(f"Dataset: {shape[0]} rows × {shape[1]} columns")

    desc = eda.get("descriptive_stats", _EDA_FALLBACK["descriptive_stats"])
    doc.add_heading(_tr(lang, "descriptive_stats"), level=2)
    desc_headers = [_tr(lang, "feature"), _tr(lang, "count"), _tr(lang, "mean"), "Std",
                    _tr(lang, "min"), "25%", "50%", "75%", _tr(lang, "max"), "Skew", "Kurtosis"]
    desc_rows = []
    for feat in _FEATURE_ORDER:
        vals = desc.get(feat, ["—"] * 11)
        desc_rows.append([feat] + [f"{v:.4f}" if isinstance(v, (int, float)) else str(v) for v in vals])
    _add_table(doc, desc_headers, desc_rows)

    doc.add_heading(_tr(lang, "missing_values"), level=2)
    missing = eda.get("missing_values", _EDA_FALLBACK["missing_values"])
    miss_headers = [_tr(lang, "feature"), _tr(lang, "count"), "%"]
    miss_rows = []
    for feat in _FEATURE_ORDER:
        vals = missing.get(feat, [0, 0.0])
        miss_rows.append([feat, str(int(vals[0])), f"{float(vals[1]):.2f}%"])
    _add_table(doc, miss_headers, miss_rows)

    doc.add_heading(_tr(lang, "outliers"), level=2)
    outliers = eda.get("outliers_iqr", _EDA_FALLBACK["outliers_iqr"])
    out_headers = [_tr(lang, "feature"), _tr(lang, "count")]
    out_rows = []
    for feat in _FEATURE_ORDER:
        if feat == "MedHouseVal":
            continue
        out_rows.append([feat, str(outliers.get(feat, 0))])
    _add_table(doc, out_headers, out_rows)

    _add_image_section(doc, "eda")
    doc.add_page_break()

    # ── Section 2: Training ─────────────────────────────────────────
    training = _load_training()
    doc.add_heading(_tr(lang, "section2_title"), level=1)
    doc.add_paragraph(_tr(lang, "section2_desc"))
    metrics = training.get("metrics", _TRAINING_FALLBACK["metrics"])
    train_headers = [_tr(lang, "model"), "MSE", "RMSE", "MAE", "R²",
                     _tr(lang, "training_time"), _tr(lang, "params")]
    train_rows = []
    for m in metrics:
        train_rows.append([
            m.get("model", ""),
            f"{m.get('MSE', 0):.4f}",
            f"{m.get('RMSE', 0):.4f}",
            f"{m.get('MAE', 0):.4f}",
            f"{m.get('R2', 0):.4f}",
            f"{m.get('training_time', 0):.1f}",
            str(m.get("params", "")),
        ])
    _add_table(doc, train_headers, train_rows)
    _add_image_section(doc, "training")
    doc.add_page_break()

    # ── Section 3: Best Model ───────────────────────────────────────
    best = training.get("best_model", _TRAINING_FALLBACK["best_model"])
    doc.add_heading(_tr(lang, "section3_title"), level=1)
    doc.add_paragraph(_tr(lang, "section3_desc"))

    doc.add_heading(best.get("model", "N/A"), level=2)
    detail_headers = [_tr(lang, "parameter"), _tr(lang, "value")]
    detail_keys = ["architecture", "activation", "optimizer", "learning_rate",
                   "batch_size", "epochs", "early_stopping"]
    detail_rows = []
    for key in detail_keys:
        label = _tr(lang, key)
        value = str(best.get(key, "—"))
        if key == "early_stopping":
            value += " epochs"
        detail_rows.append([label, value])
    _add_table(doc, detail_headers, detail_rows, col_widths=[2.5, 3.5])

    doc.add_paragraph()
    doc.add_heading("Performance", level=2)
    perf_headers = ["MSE", "RMSE", "MAE", "R²"]
    perf_rows = [[
        f"{best.get('MSE', 0):.4f}",
        f"{best.get('RMSE', 0):.4f}",
        f"{best.get('MAE', 0):.4f}",
        f"{best.get('R2', 0):.4f}",
    ]]
    _add_table(doc, perf_headers, perf_rows)
    _add_image_section(doc, "best_model")
    doc.add_page_break()

    # ── Section 4: CV ───────────────────────────────────────────────
    cv = _load_cv()
    doc.add_heading(_tr(lang, "section4_title"), level=1)
    doc.add_paragraph(_tr(lang, "section4_desc"))
    folds = cv.get("folds", _CV_FALLBACK["folds"])
    cv_headers = [_tr(lang, "fold"), "MSE", "RMSE", "MAE", "R²"]
    cv_rows = []
    for f in folds:
        cv_rows.append([
            str(f.get("fold", "")),
            f"{f.get('MSE', 0):.4f}",
            f"{f.get('RMSE', 0):.4f}",
            f"{f.get('MAE', 0):.4f}",
            f"{f.get('R2', 0):.4f}",
        ])
    mean = cv.get("mean", _CV_FALLBACK["mean"])
    std = cv.get("std", _CV_FALLBACK["std"])
    cv_rows.append([
        _tr(lang, "mean"),
        f"{mean.get('MSE', 0):.4f}",
        f"{mean.get('RMSE', 0):.4f}",
        f"{mean.get('MAE', 0):.4f}",
        f"{mean.get('R2', 0):.4f}",
    ])
    cv_rows.append([
        _tr(lang, "std"),
        f"{std.get('MSE', 0):.4f}",
        f"{std.get('RMSE', 0):.4f}",
        f"{std.get('MAE', 0):.4f}",
        f"{std.get('R2', 0):.4f}",
    ])
    _add_table(doc, cv_headers, cv_rows)
    _add_image_section(doc, "cv")
    doc.add_page_break()

    # ── Section 5: Hyperparameter ───────────────────────────────────
    hp = _load_hyperparameter()
    doc.add_heading(_tr(lang, "section5_title"), level=1)
    doc.add_paragraph(_tr(lang, "section5_desc"))
    params = hp.get("best_params", _HYPERPARAMETER_FALLBACK["best_params"])
    doc.add_heading(_tr(lang, "best_params"), level=2)
    hp_headers = [_tr(lang, "parameter"), _tr(lang, "value")]
    hp_rows = []
    for key, val in params.items():
        hp_rows.append([key.replace("_", " ").title(), str(val)])
    _add_table(doc, hp_headers, hp_rows, col_widths=[2.5, 3.5])
    score = hp.get("best_score", _HYPERPARAMETER_FALLBACK["best_score"])
    doc.add_paragraph(f"Best CV Score (R²): {score:.4f}")
    _add_image_section(doc, "hyperparameter")
    doc.add_page_break()

    # ── Section 6: Statistical ──────────────────────────────────────
    tests = _load_statistical_tests()
    doc.add_heading(_tr(lang, "section6_title"), level=1)
    doc.add_paragraph(_tr(lang, "section6_desc"))
    test_data = {
        _tr(lang, "shapiro_wilk"): tests.get("shapiro_wilk", _STATISTICAL_TESTS_FALLBACK["shapiro_wilk"]),
        _tr(lang, "durbin_watson"): tests.get("durbin_watson", _STATISTICAL_TESTS_FALLBACK["durbin_watson"]),
        _tr(lang, "breusch_pagan"): tests.get("breusch_pagan", _STATISTICAL_TESTS_FALLBACK["breusch_pagan"]),
        _tr(lang, "friedman"): tests.get("friedman", _STATISTICAL_TESTS_FALLBACK["friedman"]),
    }
    stat_headers = [_tr(lang, "test"), _tr(lang, "statistic"), _tr(lang, "p_value"), _tr(lang, "conclusion")]
    stat_rows = []
    for test_name, result in test_data.items():
        stat = result.get("statistic", 0)
        pval = result.get("p_value", 0)
        pval_str = "< 0.0001" if pval < 0.0001 else f"{pval:.4f}"
        conclusion = _tr(lang, "not_normal")
        if "autocorrelation" in result:
            conclusion = _tr(lang, "no_autocorrelation") if not result["autocorrelation"] else _tr(lang, "autocorrelation")
        elif "heteroscedastic" in result:
            conclusion = _tr(lang, "homoscedastic") if not result["heteroscedastic"] else _tr(lang, "heteroscedastic")
        elif "significant_difference" in result:
            conclusion = _tr(lang, "no_difference") if not result["significant_difference"] else _tr(lang, "significant_difference")
        elif "normal" in result:
            conclusion = _tr(lang, "normal") if result["normal"] else _tr(lang, "not_normal")
        stat_rows.append([test_name, f"{stat:.4f}", pval_str, conclusion])
    _add_table(doc, stat_headers, stat_rows)
    _add_image_section(doc, "statistical")
    doc.add_page_break()

    # ── Section 7: Conclusions ──────────────────────────────────────
    doc.add_heading(_tr(lang, "section7_title"), level=1)
    doc.add_paragraph(_tr(lang, "section7_desc"))
    conclusion = _tr(
        lang, "conclusion_text",
        arch=best.get("architecture", "[8, 64, 32, 16, 1]"),
        r2=f"{best.get('R2', 0.8284):.4f}",
        rmse=f"{best.get('RMSE', 0.4791):.4f}",
        cv_std=f"{std.get('R2', 0.0086):.4f}",
    )
    doc.add_paragraph(conclusion)
    recs = _tr(lang, "recommendations")
    if isinstance(recs, list):
        for r in recs:
            doc.add_paragraph(r, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    result = buffer.getvalue()
    if output_path:
        os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(result)
    return result

import streamlit as st
import pandas as pd
import io
import json
from datetime import datetime
from utils import api_client

if not st.session_state.get("authenticated"):
    st.warning("Please login")
    st.stop()

LANG = {
    "es": {
        "title": "Generación de Reportes",
        "subtitle": "Descargue reportes detallados del análisis completo",
        "language_selection": "Selección de Idioma",
        "download_pdf": "Descargar Reporte PDF",
        "download_word": "Descargar Reporte Word",
        "download_excel": "Descargar Reporte Excel",
        "sections": "Secciones del Reporte",
        "sections_list": [
            "Resumen Ejecutivo",
            "Análisis Exploratorio de Datos",
            "Entrenamiento de Modelos",
            "Validación Cruzada",
            "Optimización de Hiperparámetros",
            "Pruebas Estadísticas",
            "Conclusiones y Recomendaciones",
        ],
        "generating": "Generando reporte...",
        "success": "Reporte generado correctamente",
        "error": "Error al generar el reporte",
        "select_language": "Seleccione el idioma del reporte",
        "loading_data": "Cargando datos...",
        "preview": "Resumen del proyecto",
        "preview_text": "Este reporte incluye el análisis completo del dataset California Housing con 5 modelos de regresión neuronal.",
    },
    "en": {
        "title": "Report Generation",
        "subtitle": "Download detailed reports of the complete analysis",
        "language_selection": "Language Selection",
        "download_pdf": "Download PDF Report",
        "download_word": "Download Word Report",
        "download_excel": "Download Excel Report",
        "sections": "Report Sections",
        "sections_list": [
            "Executive Summary",
            "Exploratory Data Analysis",
            "Model Training",
            "Cross Validation",
            "Hyperparameter Tuning",
            "Statistical Tests",
            "Conclusions and Recommendations",
        ],
        "generating": "Generating report...",
        "success": "Report generated successfully",
        "error": "Error generating report",
        "select_language": "Select report language",
        "loading_data": "Loading data...",
        "preview": "Project Summary",
        "preview_text": "This report includes the complete analysis of the California Housing dataset with 5 neural network regression models.",
    },
}

def tr(key):
    return LANG.get(st.session_state.language, LANG["es"]).get(key, key)

lang = st.session_state.language
st.title(tr("title"))
st.markdown(f"*{tr('subtitle')}*")
st.markdown("---")

st.markdown("""
<style>
    .stDataFrame { overflow-x: auto; max-width: 100%; }
    .stDataFrame table { font-size: 12px; }
</style>
""", unsafe_allow_html=True)

report_lang = st.radio(
    tr("select_language"),
    options=["es", "en"],
    format_func=lambda x: "Español" if x == "es" else "English",
    index=0 if lang == "es" else 1,
    horizontal=True,
)

st.markdown("---")

cache = st.session_state.api_metrics_cache
if "all" not in cache:
    with st.spinner(tr("loading_data")):
        try:
            cache["all"] = api_client.get_all_metrics()
        except Exception:
            cache["all"] = {}
all_data = cache["all"]

def _build_df_sections(data):
    sections = {
        "eda": pd.DataFrame(),
        "training": pd.DataFrame(),
        "cv": pd.DataFrame(),
        "tuning": pd.DataFrame(),
        "stats": pd.DataFrame(),
    }
    eda_raw = data.get("eda", {})
    if eda_raw:
        rows = eda_raw.get("descriptive_stats", {})
        if rows:
            sections["eda"] = pd.DataFrame(rows).T.reset_index().rename(columns={"index": "Feature"})
    training_raw = data.get("training", [])
    if training_raw:
        sections["training"] = pd.DataFrame(training_raw)
    cv_raw = data.get("cross_validation", [])
    if cv_raw:
        sections["cv"] = pd.DataFrame(cv_raw)
    tuning_raw = data.get("hyperparameter_tuning", [])
    if tuning_raw:
        sections["tuning"] = pd.DataFrame(tuning_raw)
    stats_raw = data.get("statistical_tests", {})
    if stats_raw:
        rows_list = []
        for test_name, test_data in stats_raw.items():
            if isinstance(test_data, dict):
                rows_list.append({"Test": test_name, **test_data})
        if rows_list:
            sections["stats"] = pd.DataFrame(rows_list)
    return sections

# ── Generate reports ──
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("##### 📄 PDF")
    if st.button(tr("download_pdf"), use_container_width=True):
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            story.append(Paragraph("California Housing ML Report", styles["Title"]))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
            story.append(Spacer(1, 24))

            sections_data = _build_df_sections(all_data)
            for section_name, df in sections_data.items():
                if not df.empty:
                    story.append(Paragraph(section_name.upper(), styles["Heading2"]))
                    story.append(Spacer(1, 8))
                    table_data = [df.columns.tolist()] + df.astype(str).values.tolist()
                    col_w = (letter[0] - 72) / max(len(df.columns), 1)
                    t = Table(table_data, repeatRows=1, colWidths=[col_w] * len(df.columns))
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#667eea")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 7),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))

            doc.build(story)
            buf.seek(0)

            st.download_button(
                label="📥 " + tr("download_pdf"),
                data=buf,
                file_name=f"california_housing_report_{report_lang}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"{tr('error')}: {e}")

with col2:
    st.markdown("##### 📝 Word")
    if st.button(tr("download_word"), use_container_width=True):
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            buf = io.BytesIO()
            doc = Document()

            title = doc.add_heading("California Housing ML Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            sections_data = _build_df_sections(all_data)
            for section_name, df in sections_data.items():
                if not df.empty:
                    doc.add_heading(section_name.upper(), level=1)
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = "Light Grid Accent 1"
                    for j, col_name in enumerate(df.columns):
                        table.rows[0].cells[j].text = str(col_name)
                    for i, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, col_name in enumerate(df.columns):
                            row_cells[j].text = str(row[col_name])
                    doc.add_paragraph()

            doc.save(buf)
            buf.seek(0)

            st.download_button(
                label="📥 " + tr("download_word"),
                data=buf,
                file_name=f"california_housing_report_{report_lang}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"{tr('error')}: {e}")

with col3:
    st.markdown("##### 📊 Excel")
    if st.button(tr("download_excel"), use_container_width=True):
        try:
            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                sections_data = _build_df_sections(all_data)
                for section_name, df in sections_data.items():
                    if not df.empty:
                        df.to_excel(writer, sheet_name=section_name.upper(), index=False)
                meta = pd.DataFrame({
                    "Project": ["California Housing ML"],
                    "Generated": [datetime.now().strftime("%Y-%m-%d %H:%M")],
                    "Language": [report_lang],
                })
                meta.to_excel(writer, sheet_name="Summary", index=False)
            buf.seek(0)

            st.download_button(
                label="📥 " + tr("download_excel"),
                data=buf,
                file_name=f"california_housing_report_{report_lang}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"{tr('error')}: {e}")

st.markdown("---")
st.subheader(tr("preview"))
st.info(tr("preview_text"))

st.subheader(tr("sections"))
for i, section in enumerate(tr("sections_list"), 1):
    st.markdown(f"**{i}.** {section}")
